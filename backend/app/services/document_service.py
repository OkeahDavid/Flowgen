"""Document processing service: upload, chunk, embed, and search."""

import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
from io import BytesIO
from uuid import uuid4

from sqlalchemy import func, cast, String
import openai
from PyPDF2 import PdfReader
import docx

from app.models import Document
from app.database import SessionLocal

logger = logging.getLogger(__name__)


class DocumentProcessor:
    def __init__(self, openai_api_key: str):
        self.openai_client = openai.OpenAI(api_key=openai_api_key)
        self.chunk_size = 1500
        self.chunk_overlap = 200

    def extract_text(self, file_content: bytes, filename: str) -> str:
        extension = Path(filename).suffix.lower()
        if extension == ".pdf":
            pdf = PdfReader(BytesIO(file_content))
            return "\n".join(page.extract_text() for page in pdf.pages)
        elif extension in [".docx", ".doc"]:
            doc = docx.Document(BytesIO(file_content))
            return "\n".join(p.text for p in doc.paragraphs)
        elif extension in [".txt", ".md"]:
            return file_content.decode("utf-8")
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    def chunk_text(self, text: str) -> List[str]:
        words = text.split()
        chunks = []
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk = " ".join(words[i : i + self.chunk_size])
            if chunk.strip():
                chunks.append(chunk.strip())
        return chunks

    def get_embedding(self, text: str) -> List[float]:
        response = self.openai_client.embeddings.create(model="text-embedding-ada-002", input=text)
        return response.data[0].embedding

    def upload_document(self, filename: str, file_content: bytes) -> Dict[str, Any]:
        db = SessionLocal()
        try:
            text = self.extract_text(file_content, filename)
            if not text.strip():
                return {"success": False, "message": f"No text extracted from {filename}"}

            chunks = self.chunk_text(text)
            if not chunks:
                return {"success": False, "message": f"No chunks created from {filename}"}

            logger.info(f"Generating embeddings for {len(chunks)} chunks of {filename}")
            stored = 0
            for i, chunk in enumerate(chunks):
                try:
                    embedding = self.get_embedding(chunk)
                    doc = Document(
                        id=uuid4(),
                        filename=filename,
                        content=chunk,
                        content_type=Path(filename).suffix.lower().replace(".", ""),
                        file_size=len(file_content) if i == 0 else None,
                        embedding=embedding,
                        doc_metadata={"chunk_index": i, "total_chunks": len(chunks)},
                        chunk_index=i,
                        total_chunks=len(chunks),
                    )
                    db.add(doc)
                    stored += 1
                except Exception as e:
                    logger.error(f"Error processing chunk {i}: {e}")
            db.commit()
            return {"success": True, "message": f"Processed {filename}", "chunks_added": stored}
        except Exception as e:
            db.rollback()
            return {"success": False, "message": str(e)}
        finally:
            db.close()

    def search_documents(
        self, query: str, max_results: int = 5, filter_filenames: Optional[List[str]] = None
    ) -> List[Dict[str, Any]]:
        db = SessionLocal()
        try:
            query_embedding = self.get_embedding(query)
            q = db.query(Document).filter(Document.embedding.isnot(None))
            if filter_filenames:
                q = q.filter(Document.filename.in_(filter_filenames))
            documents = q.all()
            if not documents:
                return []

            from sklearn.metrics.pairwise import cosine_similarity
            import numpy as np

            doc_embeddings = np.array([doc.embedding for doc in documents])
            sims = cosine_similarity(np.array(query_embedding).reshape(1, -1), doc_embeddings)[0]
            top_idx = np.argsort(sims)[::-1][:max_results]

            results = []
            for idx in top_idx:
                if sims[idx] > 0.1:
                    doc = documents[idx]
                    results.append(
                        {
                            "filename": doc.filename,
                            "content": doc.content,
                            "relevance_score": float(sims[idx]),
                            "chunk_index": doc.chunk_index,
                            "total_chunks": doc.total_chunks,
                        }
                    )
            return results
        finally:
            db.close()

    def list_documents(self) -> Dict[str, Any]:
        db = SessionLocal()
        try:
            results = (
                db.query(
                    Document.filename,
                    func.count(Document.id).label("chunk_count"),
                    func.min(cast(Document.id, String)).label("first_id"),
                )
                .group_by(Document.filename)
                .all()
            )
            documents = [{"id": fid, "filename": fn, "chunk_count": cc} for fn, cc, fid in results]
            total = db.query(Document).count()
            return {"total_documents": len(documents), "total_chunks": total, "documents": documents}
        finally:
            db.close()

    def delete_document(self, document_id: str) -> bool:
        from uuid import UUID

        db = SessionLocal()
        try:
            doc = db.query(Document).filter(Document.id == UUID(document_id)).first()
            if not doc:
                return False
            db.query(Document).filter(Document.filename == doc.filename).delete()
            db.commit()
            return True
        except Exception:
            db.rollback()
            return False
        finally:
            db.close()


_processor_instance: Optional[DocumentProcessor] = None


def get_document_processor(openai_api_key: str) -> DocumentProcessor:
    global _processor_instance
    if _processor_instance is None:
        _processor_instance = DocumentProcessor(openai_api_key)
    return _processor_instance
