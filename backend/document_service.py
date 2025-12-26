"""
Clean document processing service with database integration.
"""
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib
from io import BytesIO
from uuid import UUID, uuid4

from sqlalchemy.orm import Session
from sqlalchemy import func, cast, String
import openai
from PyPDF2 import PdfReader
import docx

from models import Document
from database import SessionLocal

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handle document upload, chunking, embedding, and search."""
    
    def __init__(self, openai_api_key: str):
        self.openai_client = openai.OpenAI(api_key=openai_api_key)
        self.chunk_size = 1500  # words
        self.chunk_overlap = 200  # words
    
    def extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from uploaded file."""
        extension = Path(filename).suffix.lower()
        
        try:
            if extension == '.pdf':
                pdf = PdfReader(BytesIO(file_content))
                return '\n'.join(page.extract_text() for page in pdf.pages)
            
            elif extension in ['.docx', '.doc']:
                doc = docx.Document(BytesIO(file_content))
                return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
            
            elif extension in ['.txt', '.md']:
                return file_content.decode('utf-8')
            
            else:
                raise ValueError(f"Unsupported file type: {extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            raise
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk = ' '.join(words[i:i + self.chunk_size])
            if chunk.strip():
                chunks.append(chunk.strip())
        
        return chunks
    
    def get_embedding(self, text: str) -> List[float]:
        """Get OpenAI embedding for text."""
        response = self.openai_client.embeddings.create(
            model="text-embedding-ada-002",
            input=text
        )
        return response.data[0].embedding
    
    def upload_document(self, filename: str, file_content: bytes) -> Dict[str, Any]:
        """Upload and process a document."""
        db = SessionLocal()
        try:
            # Extract text
            logger.info(f"Extracting text from {filename}")
            text = self.extract_text(file_content, filename)
            
            if not text.strip():
                return {
                    "success": False,
                    "message": f"No text extracted from {filename}"
                }
            
            # Chunk text
            chunks = self.chunk_text(text)
            logger.info(f"Split {filename} into {len(chunks)} chunks")
            
            if not chunks:
                return {
                    "success": False,
                    "message": f"No chunks created from {filename}"
                }
            
            # Process each chunk
            logger.info(f"Generating embeddings for {len(chunks)} chunks")
            stored_count = 0
            
            for i, chunk in enumerate(chunks):
                try:
                    # Get embedding
                    embedding = self.get_embedding(chunk)
                    
                    # Store in database
                    doc = Document(
                        id=uuid4(),
                        filename=filename,
                        content=chunk,
                        content_type=Path(filename).suffix.lower().replace('.', ''),
                        file_size=len(file_content) if i == 0 else None,
                        embedding=embedding,
                        doc_metadata={
                            "chunk_index": i,
                            "total_chunks": len(chunks),
                            "text_length": len(chunk)
                        },
                        chunk_index=i,
                        total_chunks=len(chunks)
                    )
                    db.add(doc)
                    stored_count += 1
                    
                except Exception as e:
                    logger.error(f"Error processing chunk {i}: {e}")
                    continue
            
            db.commit()
            logger.info(f"Successfully stored {stored_count} chunks from {filename}")
            
            return {
                "success": True,
                "message": f"Processed {filename}",
                "chunks_added": stored_count
            }
            
        except Exception as e:
            db.rollback()
            logger.error(f"Error uploading {filename}: {e}")
            return {
                "success": False,
                "message": str(e)
            }
        finally:
            db.close()
    
    def search_documents(
        self, 
        query: str, 
        max_results: int = 5,
        filter_filenames: Optional[List[str]] = None
    ) -> List[Dict[str, Any]]:
        """Search documents using vector similarity."""
        db = SessionLocal()
        try:
            # Get query embedding
            logger.info(f"Searching for: {query}")
            query_embedding = self.get_embedding(query)
            
            # Build query
            q = db.query(Document)
            
            # Filter by filenames if specified
            if filter_filenames:
                q = q.filter(Document.filename.in_(filter_filenames))
            
            # Get all documents with embeddings
            documents = q.filter(Document.embedding.isnot(None)).all()
            
            if not documents:
                logger.info("No documents found")
                return []
            
            # Calculate similarities
            from sklearn.metrics.pairwise import cosine_similarity
            import numpy as np
            
            doc_embeddings = np.array([doc.embedding for doc in documents])
            query_emb = np.array(query_embedding).reshape(1, -1)
            
            similarities = cosine_similarity(query_emb, doc_embeddings)[0]
            
            # Get top results
            top_indices = np.argsort(similarities)[::-1][:max_results]
            
            results = []
            for idx in top_indices:
                if similarities[idx] > 0.1:  # Minimum threshold
                    doc = documents[idx]
                    results.append({
                        "filename": doc.filename,
                        "content": doc.content,
                        "relevance_score": float(similarities[idx]),
                        "chunk_index": doc.chunk_index,
                        "total_chunks": doc.total_chunks
                    })
            
            logger.info(f"Found {len(results)} relevant chunks")
            return results
            
        except Exception as e:
            logger.error(f"Search error: {e}")
            return []
        finally:
            db.close()
    
    def list_documents(self) -> Dict[str, Any]:
        """List all uploaded documents."""
        db = SessionLocal()
        try:
            # Get document stats grouped by filename
            results = db.query(
                Document.filename,
                func.count(Document.id).label('chunk_count'),
                func.min(cast(Document.id, String)).label('first_id')
            ).group_by(Document.filename).all()
            
            documents = [
                {
                    "id": first_id,
                    "filename": filename,
                    "chunk_count": chunk_count
                }
                for filename, chunk_count, first_id in results
            ]
            
            total_chunks = db.query(Document).count()
            
            return {
                "total_documents": len(documents),
                "total_chunks": total_chunks,
                "documents": documents
            }
            
        except Exception as e:
            logger.error(f"Error listing documents: {e}")
            return {
                "total_documents": 0,
                "total_chunks": 0,
                "documents": [],
                "error": str(e)
            }
        finally:
            db.close()
    
    def delete_document(self, document_id: str) -> bool:
        """Delete a document and all its chunks by ID."""
        db = SessionLocal()
        try:
            # Find the document by ID
            doc = db.query(Document).filter(Document.id == UUID(document_id)).first()
            
            if not doc:
                logger.warning(f"Document {document_id} not found")
                return False
            
            filename = doc.filename
            
            # Delete all chunks with this filename
            deleted = db.query(Document).filter(Document.filename == filename).delete()
            db.commit()
            
            logger.info(f"Deleted {deleted} chunks for document: {filename}")
            return True
            
        except Exception as e:
            db.rollback()
            logger.error(f"Error deleting document: {e}")
            return False
        finally:
            db.close()


# Singleton instance
_processor_instance: Optional[DocumentProcessor] = None


def get_document_processor(openai_api_key: str) -> DocumentProcessor:
    """Get or create document processor singleton."""
    global _processor_instance
    if _processor_instance is None:
        _processor_instance = DocumentProcessor(openai_api_key)
    return _processor_instance
